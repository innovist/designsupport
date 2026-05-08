"""Document parser adapter for HWP/HWPX/DOCX files.

Extracts text from Korean document formats and Office documents.
Returns None on failure (NO fake text).
"""
import os
import logging
from typing import Any

logger = logging.getLogger(__name__)

# Document parser availability flags
try:
    # Check for kordoc (HWP/HWPX parsing)
    from kordoc import HwpParser
    KORDOC_AVAILABLE = True
except ImportError:
    KORDOC_AVAILABLE = False
    HwpParser = None

try:
    # Check for python-docx (DOCX parsing)
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None


class DocParser:
    """Parse HWP/HWPX/DOCX files.

    Supports Korean document formats (HWP, HWPX) and Office documents (DOCX).
    Falls back gracefully when libraries are not available.
    """

    def __init__(self) -> None:
        """Initialize parser."""
        self._log_availability()

    def _log_availability(self) -> None:
        """Log parser availability."""
        if not KORDOC_AVAILABLE:
            logger.warning("kordoc not available. HWP/HWPX parsing disabled.")
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available. DOCX parsing disabled.")

    async def parse_file(
        self,
        file_path: str,
        file_type: str | None = None,
    ) -> str | None:
        """Parse document file and extract text.

        Args:
            file_path: Path to document file
            file_type: Optional file type hint (hwp, hwpx, docx)

        Returns:
            Extracted text content or None if parsing fails

        Raises:
            ValueError: If file type is not supported
        """
        if not os.path.exists(file_path):
            logger.error(f"File not found: {file_path}")
            return None

        # Detect file type if not provided
        if file_type is None:
            file_type = self._detect_ext(file_path)

        # Route to appropriate parser
        if file_type in ["hwp", "hwpx"]:
            return await self._parse_kordoc(file_path, file_type)
        elif file_type == "docx":
            return await self._parse_docx(file_path)
        else:
            logger.error(f"Unsupported file type: {file_type}")
            return None

    async def _parse_kordoc(self, file_path: str, file_type: str) -> str | None:
        """Parse HWP/HWPX file using kordoc.

        Args:
            file_path: Path to HWP/HWPX file
            file_type: File type (hwp or hwpx)

        Returns:
            Extracted text or None
        """
        if not KORDOC_AVAILABLE:
            logger.error(f"kordoc not available for {file_type} parsing")
            return None

        try:
            parser = HwpParser()
            text = parser.parse_file(file_path)

            if text and text.strip():
                logger.info(f"Extracted {len(text)} characters from {file_path}")
                return text
            else:
                logger.warning(f"No text extracted from {file_path}")
                return None

        except Exception as e:
            logger.error(f"HWP/HWPX parsing failed for {file_path}: {e}")
            return None

    async def _parse_docx(self, file_path: str) -> str | None:
        """Parse DOCX file using python-docx.

        Args:
            file_path: Path to DOCX file

        Returns:
            Extracted text or None
        """
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available for DOCX parsing")
            return None

        try:
            doc = Document(file_path)

            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text and paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))

            if not text_parts:
                logger.warning(f"No text extracted from DOCX: {file_path}")
                return None

            full_text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(full_text)} characters from {file_path}")
            return full_text

        except Exception as e:
            logger.error(f"DOCX parsing failed for {file_path}: {e}")
            return None

    async def detect_type(self, file_path: str) -> str | None:
        """Detect document file type.

        Args:
            file_path: Path to file

        Returns:
            Detected file type or None
        """
        if not os.path.exists(file_path):
            return None

        return self._detect_ext(file_path)

    def _detect_ext(self, file_path: str) -> str | None:
        """Detect file type from extension.

        Args:
            file_path: Path to file

        Returns:
            File type or None
        """
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        if ext == ".hwp":
            return "hwp"
        elif ext == ".hwpx":
            return "hwpx"
        elif ext == ".docx":
            return "docx"

        return None


def get_doc_parser() -> DocParser:
    """Get singleton DocParser instance.

    Returns:
        DocParser instance
    """
    return DocParser()
